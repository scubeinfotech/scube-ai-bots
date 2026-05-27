"""
Document parser — extracts plain text from PDF, DOCX, and text files.
Used by the document upload endpoint before vector indexing.
"""
from io import BytesIO


def parse_document(filename: str, content_bytes: bytes) -> str:
    """Extract plain text from file bytes based on file extension."""
    ext = (filename or "").rsplit(".", 1)[-1].lower() if "." in (filename or "") else "txt"

    if ext == "pdf":
        return _parse_pdf(content_bytes)
    elif ext in ("docx", "doc"):
        return _parse_docx(content_bytes)
    else:
        # Plain text, markdown, CSV, etc.
        return content_bytes.decode("utf-8", errors="ignore")


def _parse_pdf(content_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(BytesIO(content_bytes))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                pages.append(text.strip())
        return "\n\n".join(pages)
    except ImportError:
        return content_bytes.decode("utf-8", errors="ignore")
    except Exception as exc:
        return f"[PDF extraction failed: {exc}]"


def _parse_docx(content_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(BytesIO(content_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))
        return "\n".join(paragraphs)
    except ImportError:
        return content_bytes.decode("utf-8", errors="ignore")
    except Exception as exc:
        return f"[DOCX extraction failed: {exc}]"
